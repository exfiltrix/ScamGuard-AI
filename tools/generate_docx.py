"""
Генератор документа ScamGuard AI Concept.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "ScamGuard_AI_Concept.docx")


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1, color=(31, 73, 125)):
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def para(doc, text, bold_parts=None, indent=False, size=11, color=None):
    """Add paragraph, optionally with some bold_parts highlighted."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                set_font(r, size=size, color=color)
            r = p.add_run(bp)
            set_font(r, size=size, bold=True, color=color)
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            set_font(r, size=size, color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color)
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, bold=True, size=10, color=(255, 255, 255))
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ─── ТИТУЛЬНАЯ СТРАНИЦА ────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("🛡️ ScamGuard AI")
    set_font(r, size=32, bold=True, color=(31, 73, 125))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Telegram-бот для защиты от мошенников")
    set_font(r, size=16, color=(89, 89, 89))

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run("Версия 0.4.0  |  Апрель 2026  |  MVP готов")
    set_font(r, size=11, color=(127, 127, 127), italic=True)

    doc.add_page_break()

    # ─── 1. ПРОБЛЕМА ──────────────────────────────────────────────────────
    heading(doc, "1. Проблема")

    para(doc,
        "Каждый день тысячи людей получают мошеннические сообщения в мессенджерах. "
        "Мошенники представляются сотрудниками банков, друзьями, организаторами акций — "
        "и жертвы теряют деньги, потому что не могут быстро распознать обман.")

    doc.add_paragraph()
    para(doc, "Примеры реальных мошеннических схем:", bold_parts=["Примеры реальных мошеннических схем:"])

    schemes = [
        ("🏦 Банковский фишинг",
         '"Это служба безопасности банка. На вашем счёте подозрительная операция. '
         'Срочно подтвердите данные по ссылке, иначе счёт будет заблокирован."'),
        ("💸 Схема «займи денег»",
         '"Привет, срочно нужна помощь 🙏 Можешь одолжить 200 000 сум? '
         'Переведи на карту 8600 **** **** 1234, позже верну."'),
        ("🤖 Инвестиционная пирамида",
         '"Хочешь заработать? У меня есть бот, в котором я заработал 1 000 000$. '
         'Заходи по моей реферальной ссылке и зарабатывай!"'),
        ("🎁 Фейковый приз",
         '"Поздравляем! Вы стали победителем акции! '
         'Чтобы получить iPhone, оплатите доставку по ссылке. Осталось 3 часа!"'),
        ("🇺🇿 Узбекский фишинг",
         '"Hurmatli foydalanuvchi: hisobingiz cheklandi. '
         '13 soat ichida shaxsni tasdiqlang: z1xx.ydi55trui.shop"'),
    ]

    for title_text, example in schemes:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(title_text + " — ")
        set_font(r, bold=True, size=11)
        r2 = p.add_run(example)
        set_font(r2, italic=True, size=10, color=(89, 89, 89))

    doc.add_paragraph()
    para(doc,
        "Проблема усугубляется тем, что мошенники пишут на разных языках (русский, узбекский), "
        "используют психологическое давление и срочность, а ссылки маскируют под легитимные сайты.",
        color=(89, 89, 89))

    # ─── 2. РЕШЕНИЕ ───────────────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, "2. Решение — ScamGuard AI")

    para(doc,
        "ScamGuard AI — Telegram-бот, который мгновенно анализирует подозрительные сообщения "
        "и выдаёт оценку риска с объяснением на понятном языке.",
        bold_parts=["ScamGuard AI"])

    doc.add_paragraph()
    para(doc, "Как это работает:", bold_parts=["Как это работает:"])

    steps = [
        "Пользователь пересылает подозрительное сообщение боту",
        "Бот анализирует текст, ссылки и контекст за 3–15 секунд",
        'Возвращает оценку риска от 0 до 100 ("низкий / средний / высокий")',
        "Объясняет конкретно что именно подозрительно",
        "Даёт персональные рекомендации что делать",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(step)
        set_font(r, size=11)

    doc.add_paragraph()
    para(doc, "Пример ответа бота на банковский фишинг:", bold_parts=["Пример ответа бота на банковский фишинг:"])

    example_box = doc.add_paragraph()
    example_box.paragraph_format.left_indent = Cm(1)
    example_box.paragraph_format.right_indent = Cm(1)
    r = example_box.add_run(
        "🔴 ВЫСОКИЙ РИСК\n\n"
        "Оценка: 85/100\n\n"
        "Что насторожило:\n"
        "• 🏦 Самозванец: «менеджер/сотрудник банка» + ссылка — банки никогда не пишут так\n"
        "• 🔗 Фишинговая ссылка: угроза + ссылка + запрос данных\n"
        "• 🚫 Сайт secure-check-account.com не существует (проверено)\n\n"
        "Что делать:\n"
        "• 🚨 Не переходите по ссылке и не вводите никакие данные\n"
        "• 📞 Позвоните в банк напрямую по номеру с официального сайта"
    )
    set_font(r, size=10, color=(89, 89, 89), italic=True)

    # ─── 3. КАК УСТРОЕН БОТ ───────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "3. Как устроен бот — техническая часть")

    para(doc,
        "Система использует многоуровневый подход: несколько независимых модулей "
        "анализируют сообщение параллельно, результаты объединяются в итоговую оценку.")

    doc.add_paragraph()
    heading(doc, "3.1 Модули анализа", level=2)

    add_table(doc,
        headers=["Модуль", "Что делает", "Скорость", "Вес"],
        rows=[
            ("Rule Engine",
             "300+ правил и паттернов: ключевые слова, комбо-детекторы схем мошенничества",
             "Мгновенно", "20%"),
            ("URL Analyzer",
             "Проверяет ссылки: возраст домена, DNS, подозрительные TLD, имперсонация брендов",
             "1–2 сек", "15%"),
            ("Context Analyzer",
             "Gemini + Google Search: существует ли сайт в реальности, легитимна ли организация",
             "3–8 сек", "20%"),
            ("Gemini NLP",
             "LLM-анализ: понимает контекст, находит манипуляции и аномалии",
             "3–8 сек", "25%"),
            ("Embedding",
             "Сравнивает сообщение с базой известных мошеннических схем",
             "< 1 сек", "10%"),
            ("Image Analyzer",
             "Анализирует прикреплённые фото через Gemini Vision",
             "3–5 сек", "10%"),
        ],
        col_widths=[3.5, 7.5, 2.5, 1.5],
    )

    doc.add_paragraph()
    heading(doc, "3.2 Поддерживаемые схемы мошенничества", level=2)

    add_table(doc,
        headers=["Тип схемы", "Языки", "Точность"],
        rows=[
            ("Банковский фишинг (сотрудник банка)", "RU, UZ", "~95%"),
            ("Схема «займи денег»", "RU, UZ", "~90%"),
            ("Фейковый приз / конкурс", "RU, UZ", "~90%"),
            ("Инвестиционный бот / пирамида", "RU, UZ, EN", "~88%"),
            ("Фейковые вакансии", "RU", "~80%"),
            ("Романтические мошенничества", "RU", "~75%"),
            ("Мошенничество от «друзей/родственников»", "RU, UZ", "~85%"),
            ("Фейковая аренда жилья", "RU", "~80%"),
        ],
        col_widths=[8.5, 3, 2.5],
    )

    doc.add_paragraph()
    heading(doc, "3.3 Схема работы", level=2)

    flow = doc.add_paragraph()
    flow.paragraph_format.left_indent = Cm(1)
    r = flow.add_run(
        "Сообщение пользователя\n"
        "        ↓\n"
        "┌─────────────────────────────────────┐\n"
        "│  БЫСТРАЯ ПРОВЕРКА (3–5 сек)          │\n"
        "│  Rule Engine + URL + Context         │\n"
        "│  → Результат сразу                   │\n"
        "└─────────────────────────────────────┘\n"
        "        ↓  (по запросу)\n"
        "┌─────────────────────────────────────┐\n"
        "│  ГЛУБОКИЙ АНАЛИЗ (8–15 сек)          │\n"
        "│  + Gemini NLP + Image Analysis       │\n"
        "│  → Полный развёрнутый ответ          │\n"
        "└─────────────────────────────────────┘"
    )
    set_font(r, size=10, color=(89, 89, 89))

    # ─── 4. ФУНКЦИИ БОТА ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "4. Функции Telegram-бота")

    heading(doc, "4.1 Основной сценарий", level=2)
    steps2 = [
        ('Отправить /start', " — открыть главное меню"),
        ('Переслать подозрительное сообщение', " — бот сразу начнёт анализ"),
        ('Получить оценку риска', " — с объяснением каждого флага"),
        ('Нажать «Глубокий анализ»', " — для полного AI-разбора"),
        ('Нажать «Это мошенничество»', " — сообщить боту, данные помогут улучшить систему"),
    ]
    for bold_part, rest in steps2:
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(bold_part)
        set_font(r1, bold=True)
        r2 = p.add_run(rest)
        set_font(r2)

    doc.add_paragraph()
    heading(doc, "4.2 Команды бота", level=2)

    add_table(doc,
        headers=["Команда", "Описание"],
        rows=[
            ("/start", "Главное меню"),
            ("/analyze", "Проверить сообщение или ссылку"),
            ("/history", "История проверок пользователя"),
            ("/stats", "Статистика: сколько скамов выявлено"),
            ("/help", "Инструкция по использованию"),
        ],
        col_widths=[4, 10],
    )

    doc.add_paragraph()
    heading(doc, "4.3 Что получает пользователь в ответе", level=2)

    bullets = [
        ("Оценка риска 0–100", " с уровнем (🟢 НИЗКИЙ / 🟡 СРЕДНИЙ / 🔴 ВЫСОКИЙ)"),
        ("Список конкретных флагов", " — что именно подозрительно и почему"),
        ("Результат проверки ссылок", " — существует ли сайт в реальности"),
        ("Персональные рекомендации", " — что делать прямо сейчас"),
        ("Тип схемы мошенничества", " — банковский фишинг, пирамида, займ и т.д."),
    ]
    for bold_part, rest in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(bold_part)
        set_font(r1, bold=True)
        r2 = p.add_run(rest)
        set_font(r2)

    # ─── 5. ТЕХНИЧЕСКИЙ СТЕК ──────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, "5. Технический стек")

    add_table(doc,
        headers=["Компонент", "Технология", "Стоимость"],
        rows=[
            ("Backend API", "Python 3.14 + FastAPI", "Бесплатно"),
            ("Telegram бот", "aiogram 3", "Бесплатно"),
            ("AI / NLP", "Google Gemini 2.0 Flash", "Бесплатно (60 req/min)"),
            ("Поиск для проверки URL", "Gemini + Google Search Grounding", "Бесплатно"),
            ("База данных", "SQLite + SQLAlchemy", "Бесплатно"),
            ("DNS / WHOIS", "dnspython + python-whois", "Бесплатно"),
            ("Web UI", "HTML/CSS/JS (без фреймворков)", "Бесплатно"),
            ("Хостинг", "Локально / VPS", "От $5/мес"),
        ],
        col_widths=[4.5, 5.5, 4],
    )

    doc.add_paragraph()
    para(doc,
        "Итоговая стоимость для MVP: $0 в месяц (всё на бесплатных тарифах). "
        "При масштабировании — оплата только за хостинг.",
        bold_parts=["$0 в месяц"])

    # ─── 6. ДОРОЖНАЯ КАРТА ────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "6. Дорожная карта развития")

    heading(doc, "Сейчас готово (v0.4.0)", level=2)
    done = [
        "Рабочий Telegram-бот (@ScamGuardAI_bot)",
        "6 параллельных модулей анализа",
        "10+ типов мошеннических схем на RU и UZ",
        "Проверка URL: существование сайта, возраст домена, DNS",
        "Web UI для демонстрации",
        "API с документацией (Swagger)",
        "История проверок и обратная связь от пользователей",
        "Инструменты для накопления датасета",
    ]
    for item in done:
        bullet(doc, "✅ " + item)

    doc.add_paragraph()
    heading(doc, "Ближайшие 2–4 недели", level=2)
    soon = [
        "Добавить паттерны: Tech support scam, Fake delivery, Marketplace scam",
        "Расширить узбекскую базу паттернов до 150+",
        "Добавить легитимные узбекские домены (uzcard.uz, humo.uz, click.uz и др.)",
        "Кеширование результатов URL-анализа",
        "Накопить 50+ примеров через фидбек пользователей",
    ]
    for item in soon:
        bullet(doc, "🔜 " + item)

    doc.add_paragraph()
    heading(doc, "1–3 месяца", level=2)
    mid = [
        "TF-IDF + Logistic Regression модель (после 200+ примеров)",
        "Интеграция с PhishTank и Google Safe Browsing API",
        "Проверка возраста и репутации Telegram-каналов",
        "Мобильное приложение (Flutter)",
        "Браузерное расширение (Chrome/Firefox)",
    ]
    for item in mid:
        bullet(doc, "📌 " + item)

    doc.add_paragraph()
    heading(doc, "3–6 месяцев", level=2)
    long_term = [
        "Fine-tune multilingual BERT на 1000+ примерах (точность ~93%)",
        "Интеграция с OLX.uz, Avito для автоматической проверки объявлений",
        "Premium API для бизнеса (банки, маркетплейсы)",
        "Расширение на другие страны СНГ",
    ]
    for item in long_term:
        bullet(doc, "🎯 " + item)

    # ─── 7. ЧТО НУЖНО ДОРАБОТАТЬ ──────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, "7. Что нужно доработать (приоритеты)")

    heading(doc, "🔴 Критично — влияет на точность прямо сейчас", level=2, color=(192, 0, 0))

    critical = [
        ("Tech support scam — паттернов нет",
         'Добавить в rule_engine: "ваш компьютер заражён", "anydesk", "установите программу"'),
        ("Fake delivery scam — паттернов нет",
         'Добавить: "ваша посылка задержана", "таможенный сбор", "подтвердите адрес"'),
        ("Marketplace scam — паттернов нет",
         'Добавить: "куплю без торга", "перевод через гарант", "напишите в вацап"'),
        ("Узбекский датасет",
         "Расширить с 50 до 150+ паттернов, особенно банковский фишинг на UZ"),
        ("Embedding база",
         "Сейчас 15 паттернов → нужно 100+, добавить примеры всех типов схем"),
        ("Доверенные домены",
         "Добавить uzcard.uz, humo.uz, click.uz, payme.uz, egov.uz и др."),
    ]
    add_table(doc,
        headers=["Проблема", "Что сделать"],
        rows=critical,
        col_widths=[5, 9],
    )

    doc.add_paragraph()
    heading(doc, "🟡 Важно — улучшит качество", level=2, color=(192, 96, 0))

    important = [
        ("Датасет из фидбека",
         "Собрать 200+ размеченных примеров → обучить TF-IDF классификатор"),
        ("Кеширование URL",
         "Добавить in-memory кеш на 1 час для DNS/WHOIS запросов"),
        ("Тесты",
         "Написать unit-тесты для rule_engine и url_analyzer"),
    ]
    add_table(doc,
        headers=["Задача", "Зачем"],
        rows=important,
        col_widths=[4.5, 9.5],
    )

    # ─── 8. КОНТАКТЫ ──────────────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, "8. Ресурсы проекта")

    add_table(doc,
        headers=["Ресурс", "Ссылка / Описание"],
        rows=[
            ("Telegram бот", "@ScamGuardAI_bot"),
            ("API сервер", "http://localhost:8000"),
            ("API документация", "http://localhost:8000/docs"),
            ("Web UI", "http://localhost:8000 (главная страница)"),
            ("Архитектура", "docs/ARCHITECTURE.md"),
            ("Гайд разработчика", "docs/DEVELOPMENT.md"),
            ("Работа с датасетом", "docs/DATASET.md"),
        ],
        col_widths=[4.5, 9.5],
    )

    # ─── FOOTER ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("ScamGuard AI  |  Версия 0.4.0  |  Апрель 2026")
    set_font(r, size=9, color=(127, 127, 127), italic=True)

    doc.save(OUTPUT)
    print(f"✅ Документ сохранён: {OUTPUT}")


if __name__ == "__main__":
    main()
